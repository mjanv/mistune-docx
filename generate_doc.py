# -*- coding: utf-8 -*-
import os
import shutil
import glob
import re
import itertools
import argparse

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
import mistune


class MathBlockGrammar(mistune.BlockGrammar):
    block_math = re.compile(r"^\$\$(.*?)\$\$", re.DOTALL)


class MathBlockLexer(mistune.BlockLexer):
    default_rules = ['block_math'] + mistune.BlockLexer.default_rules

    def __init__(self, rules=None, **kwargs):
        if rules is None:
            rules = MathBlockGrammar()
        super(MathBlockLexer, self).__init__(rules, **kwargs)

    def parse_block_math(self, m):
        """Parse a $$math$$ block"""
        self.tokens.append({'type': 'block_math', 'text': m.group(1)})


class MarkdownWithMath(mistune.Markdown):
    def __init__(self, renderer, **kwargs):
        kwargs['block'] = MathBlockLexer
        super(MarkdownWithMath, self).__init__(renderer, **kwargs)

    def output_block_math(self):
        return self.renderer.block_math(self.token['text'])


class PythonDocxRenderer(mistune.Renderer):
    def __init__(self, **kwds):
        super(PythonDocxRenderer, self).__init__(**kwds)
        self.table_memory = []
        self.img_counter = 0

    def header(self, text, level, raw):
        return "p = document.add_heading('', %d)\n" % (level - 1) + text

    def paragraph(self, text):
        if 'add_picture' in text:
            return text
        add_break = '' if text.endswith(':")\n') else 'p.add_run().add_break()'
        return '\n'.join(('p = document.add_paragraph()', text, add_break)) + '\n'

    def list(self, body, ordered):
        return body + '\np.add_run().add_break()\n'

    def list_item(self, text):
        return '\n'.join(("p = document.add_paragraph('', style = 'BasicUserList')", text))

    def table(self, header, body):
        number_cols = header.count('\n') - 2
        number_rows = int(len(self.table_memory) / number_cols)
        cells = ["table.rows[%d].cells[%d].paragraphs[0]%s\n" % (i, j, self.table_memory.pop(0)[1:]) for i, j in itertools.product(range(number_rows), range(number_cols))]
        return '\n'.join(["table = document.add_table(rows=%d, cols=%d, style = 'BasicUserTable')" % (number_rows, number_cols)] + cells) + 'document.add_paragraph().add_run().add_break()\n'

    def table_cell(self, content, **flags):
        self.table_memory.append(content)
        return content

    # SPAN LEVEL
    def text(self, text):
        return "p.add_run(\"%s\")\n" % text

    def emphasis(self, text):
        return text[:-1] + '.italic = True\n'

    def double_emphasis(self, text):
        return text[:-1] + '.bold = True\n'

    def block_code(self, code, language):
        code = code.replace('\n', '\\n')
        return "p = document.add_paragraph()\np.add_run(\"%s\")\np.style = 'BasicUserQuote'\np.add_run().add_break()\n" % code

    def link(self, link, title, content):
        return "%s (%s)" % (content, link)

    def image(self, src, title, alt_text):
        return '\n'.join((
            "p = document.add_paragraph()",
            "p.alignment = WD_ALIGN_PARAGRAPH.CENTER",
            "p.space_after = Pt(18)",
            "run = p.add_run()",
            "run.add_picture(\'%s\')" % src if "tmp" in src else "run.add_picture(\'%s\', width=Cm(15))" % src,
            "run.add_break()",
            "run.add_text(\'%s\')" % alt_text,
            "run.font.italic = True",
            "run.add_break()"
            )) + '\n'

    def hrule(self):
        return "document.add_page_break()\n"

    def block_math(self, text):
        import sympy
        if not os.path.exists('tmp'):
            os.makedirs('tmp')
        filename = 'tmp/tmp%d.png' % self.img_counter
        self.img_counter = self.img_counter + 1
        sympy.preview(r'$$%s$$' % text, output='png', viewer='file', filename=filename, euler=False)
        return self.image(filename, None, "Equation " + str(self.img_counter - 1))

parser = argparse.ArgumentParser(description='Generate Docx reports using a Docx reference template and Markdown files')
parser.add_argument('output', default=None, help='Output file')
parser.add_argument('--template', default=None, help='Docx template')
parser.add_argument('--files', default="*.md", help='Regex for Markdown files')
args = parser.parse_args()

document = Document(os.path.abspath(args.template)) if args.template else Document()

T = []

for part in sorted(glob.glob(args.files)):
    with open(part, 'r', encoding="utf-8") as f:
        T.append(f.read())

renderer = PythonDocxRenderer()

exec(MarkdownWithMath(renderer=renderer)('\n'.join(T)))
document.save(os.path.abspath(args.output))
if os.path.exists('tmp'):
    shutil.rmtree('tmp')
